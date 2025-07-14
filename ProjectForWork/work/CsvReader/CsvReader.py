import pandas as pd
from docx import Document
from typing import Optional
from dateutil import parser


class ReadAndWrite:
    """
    Класс для чтения CSV-файлов и создания документов Word.
    """

    def __init__(self) -> None:
        self.content: Optional[pd.DataFrame] = None
        self.csvFile: Optional[str] = None
        self.table = None
        self.doc = None

    def set_value(self) -> Optional[pd.DataFrame]:
        """
        Устанавливает путь к CSV-файлу и читает данные.
        """
        self.csvFile = input('Введите путь к CSV-файлу: ').strip()
        try:
            self.content = pd.read_csv(self.csvFile, on_bad_lines='skip')
            if self.content.empty:
                print("Файл пустой.")
                return None
            return self.content
        except Exception as e:
            print(f"Ошибка чтения файла: {e}")
            return None

    def create_document(self, columns: list[str]) -> None:
        """
        Создает Word-документ с заголовком и таблицей.
        """
        self.doc = Document()
        self.doc.add_heading('Таблица из CSV', 0)

        self.table = self.doc.add_table(rows=1, cols=len(columns))
        hdr_cells = self.table.rows[0].cells
        for i, column in enumerate(columns):
            hdr_cells[i].text = str(column)

    def add_info(self, columns: list[str]) -> None:
        """
        Добавляет строки данных в таблицу документа.
        Форматирует даты, если возможно, иначе записывает как строку.
        """
        for _, row in self.content.iterrows():
            row_cells = self.table.add_row().cells
            for i, column in enumerate(columns):
                value = row[column]
                text = ""

                if pd.isnull(value):
                    text = ""
                else:
                    try:
                        # Попытка распознать как дату
                        parsed_date = parser.parse(str(value), dayfirst=True)
                        text = parsed_date.strftime('%d.%m.%Y')
                    except (ValueError, TypeError):
                        # Если не дата — просто строка
                        text = str(value)

                row_cells[i].text = text

        self.doc.save('output.docx')

    def check_date_in_columns(self, columns: list[str]) -> None:
        """
        Проверка наличия дат в указанных столбцах.
        """
        for column in columns:
            if column in self.content.columns:
                count = 0
                for val in self.content[column].dropna():
                    try:
                        parser.parse(str(val), dayfirst=True)
                        count += 1
                    except Exception:
                        continue
                print(f"Столбец '{column}': найдено {count} значений, похожих на дату.")
            else:
                print(f"Столбец '{column}' не найден в данных.")


def main():
    """
    Точка входа скрипта.
    """
    processor = ReadAndWrite()
    processor.content = processor.set_value()

    if processor.content is not None:
        columns = ['Fee Earner Name', 'Date']  # Укажи нужные столбцы
        columns = [col.strip() for col in columns if col.strip()]

        if all(col in processor.content.columns for col in columns):
            processor.check_date_in_columns(columns)
            processor.create_document(columns)
            processor.add_info(columns)
            print("Документ успешно создан: output.docx")
        else:
            print("Один или несколько указанных столбцов не найдены в CSV-файле.")


if __name__ == '__main__':
    main()
